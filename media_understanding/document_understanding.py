"""
文档理解模块

提供文档内容提取和分析功能
"""

import asyncio
import io
import json
import time
from typing import Dict, List, Optional, Any, Tuple
from pathlib import Path
import base64

# PDF处理
try:
    import PyPDF2
    from pdfplumber import PDF
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# Office文档处理
try:
    import docx
    import pandas as pd
    from openpyxl import load_workbook
    EXCEL_AVAILABLE = True
    EXCEL_ERROR = None
except ImportError as e:
    EXCEL_AVAILABLE = False
    EXCEL_ERROR = str(e)

# 图像处理（用于OCR）
try:
    from PIL import Image
    import pytesseract
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

# 文本处理
import re
import csv
from html.parser import HTMLParser


from .types import (
    DocumentExtractionRequest, 
    DocumentExtractionResult,
    MediaUnderstandingProvider,
    MediaUnderstandingCapability,
)


class BaseDocumentUnderstandingProvider(MediaUnderstandingProvider):
    """文档理解Provider基类"""
    
    def __init__(self, provider_id: str):
        super().__init__(provider_id)
    
    def get_capabilities(self) -> List[MediaUnderstandingCapability]:
        """返回支持的媒体理解能力"""
        return [MediaUnderstandingCapability.DOCUMENT]
    
    async def transcribe_audio(self, request):
        """音频转录功能由子类实现"""
        raise NotImplementedError(f"{self.id} provider does not support audio transcription")
    
    async def describe_image(self, request):
        """图像描述功能由子类实现"""
        raise NotImplementedError(f"{self.id} provider does not support image description")
    
    async def describe_video(self, request):
        """视频描述功能由子类实现"""
        raise NotImplementedError(f"{self.id} provider does not support video description")


class PDFProvider(BaseDocumentUnderstandingProvider):
    """PDF文档理解Provider"""
    
    def __init__(self):
        super().__init__("pdf")
        self.available = PDF_AVAILABLE
    
    async def extract_content(self, request: DocumentExtractionRequest) -> DocumentExtractionResult:
        """提取PDF内容"""
        if not self.available:
            raise Exception("PDF processing libraries not available")
        
        try:
            pdf_buffer = io.BytesIO(request.buffer)
            
            # 使用pdfplumber提取内容（更好的表格和文本提取）
            with PDF(pdf_buffer) as pdf:
                pages = len(pdf.pages)
                
                # 限制页面数量
                max_pages = request.max_pages or pages
                pages_to_process = min(pages, max_pages)
                
                text_content = []
                tables_content = []
                
                for i in range(pages_to_process):
                    page = pdf.pages[i]
                    
                    # 提取文本
                    if request.extract_text:
                        page_text = page.extract_text()
                        if page_text:
                            text_content.append(f"--- 第{i+1}页 ---\n{page_text}")
                    
                    # 提取表格
                    if request.extract_tables:
                        page_tables = page.extract_tables()
                        if page_tables:
                            for j, table in enumerate(page_tables):
                                table_text = f"表格{i+1}-{j+1}:\n"
                                table_text += "\n".join(["\t".join([str(cell) if cell else "" for cell in row]) for row in table])
                                tables_content.append(table_text)
                
                # 合并文本内容
                full_text = "\n\n".join(text_content)
                
                # 提取图像（如果支持且需要）
                extracted_images = []
                if request.extract_images and OCR_AVAILABLE:
                    extracted_images = await self._extract_images_from_pdf(pdf, pages_to_process)
                
                return DocumentExtractionResult(
                    text=full_text,
                    pages=pages_to_process,
                    tables=tables_content if request.extract_tables else None,
                    images=extracted_images if request.extract_images else None,
                    metadata={
                        "total_pages": pages,
                        "processed_pages": pages_to_process,
                        "extraction_method": "pdfplumber"
                    }
                )
                
        except Exception as e:
            raise Exception(f"PDF extraction failed: {str(e)}")
    
    async def _extract_images_from_pdf(self, pdf, max_pages: int) -> List[bytes]:
        """从PDF中提取图像（基础实现）"""
        # pdfplumber的图像提取功能有限
        # 这里提供基础实现
        images = []
        
        try:
            for page_num in range(min(max_pages, 3)):  # 最多处理3页
                page = pdf.pages[page_num]
                # 这里需要更复杂的图像提取逻辑
                # 临时返回空列表
                pass
        except Exception:
            pass
        
        return images


class OfficeDocumentProvider(BaseDocumentUnderstandingProvider):
    """Office文档理解Provider"""
    
    def __init__(self):
        super().__init__("office")
        self.available = EXCEL_AVAILABLE
        self.excel_error = EXCEL_ERROR
    
    async def extract_content(self, request: DocumentExtractionRequest) -> DocumentExtractionResult:
        """提取Office文档内容"""
        if not self.available:
            raise Exception(f"Excel processing not available: {self.excel_error}")
        
        try:
            doc_buffer = io.BytesIO(request.buffer)
            file_name = request.file_name.lower()
            
            # 根据文件类型选择处理方法
            if file_name.endswith('.docx') or file_name.endswith('.doc'):
                return await self._extract_word_content(doc_buffer)
            elif file_name.endswith('.xlsx') or file_name.endswith('.xls'):
                return await self._extract_excel_content(doc_buffer)
            elif file_name.endswith('.pptx') or file_name.endswith('.ppt'):
                return await self._extract_powerpoint_content(doc_buffer)
            else:
                raise Exception(f"Unsupported office document type: {file_name}")
                
        except Exception as e:
            raise Exception(f"Office document extraction failed: {str(e)}")
    
    async def _extract_word_content(self, doc_buffer: io.BytesIO) -> DocumentExtractionResult:
        """提取Word文档内容"""
        try:
            doc = docx.Document(doc_buffer)
            
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            text_content = "\n".join(paragraphs)
            
            # 提取表格
            tables_content = []
            for i, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_data.append(row_data)
                
                if table_data:
                    table_text = f"表格{i+1}:\n"
                    table_text += "\n".join(["\t".join(row) for row in table_data])
                    tables_content.append(table_text)
            
            return DocumentExtractionResult(
                text=text_content,
                tables=tables_content if tables_content else None,
                metadata={
                    "document_type": "word",
                    "paragraphs_count": len(paragraphs),
                    "tables_count": len(tables_content)
                }
            )
            
        except Exception as e:
            raise Exception(f"Word document extraction failed: {str(e)}")
    
    async def _extract_excel_content(self, doc_buffer: io.BytesIO) -> DocumentExtractionResult:
        """提取Excel文档内容"""
        try:
            # 使用openpyxl处理.xlsx文件
            workbook = load_workbook(doc_buffer, read_only=True, data_only=True)
            
            all_sheets_data = []
            total_tables = 0
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                
                # 获取工作表数据
                sheet_data = []
                for row in sheet.iter_rows(values_only=True):
                    row_data = [str(cell) if cell is not None else "" for cell in row]
                    if any(row_data):  # 跳过空行
                        sheet_data.append(row_data)
                
                if sheet_data:
                    all_sheets_data.append(f"--- 工作表: {sheet_name} ---\n")
                    all_sheets_data.append("\n".join(["\t".join(row) for row in sheet_data]))
                    all_sheets_data.append("\n\n")
                    total_tables += 1
            
            text_content = "".join(all_sheets_data)
            
            return DocumentExtractionResult(
                text=text_content,
                metadata={
                    "document_type": "excel",
                    "sheets_count": len(workbook.sheetnames),
                    "tables_count": total_tables
                }
            )
            
        except Exception as e:
            raise Exception(f"Excel document extraction failed: {str(e)}")
    
    async def _extract_powerpoint_content(self, doc_buffer: io.BytesIO) -> DocumentExtractionResult:
        """提取PowerPoint文档内容"""
        # PowerPoint处理比较复杂，这里提供基础实现
        try:
            doc = docx.Document(doc_buffer)
            
            # 提取幻灯片标题和内容
            slides_content = []
            slide_num = 1
            
            for para in doc.paragraphs:
                if para.text.strip():
                    if len(slides_content) == 0 or para.style.name.startswith('Heading'):
                        slides_content.append(f"\n--- 幻灯片 {slide_num} ---")
                        slide_num += 1
                    slides_content.append(para.text)
            
            text_content = "\n".join(slides_content)
            
            return DocumentExtractionResult(
                text=text_content,
                metadata={
                    "document_type": "powerpoint",
                    "slides_count": slide_num - 1
                }
            )
            
        except Exception as e:
            raise Exception(f"PowerPoint document extraction failed: {str(e)}")


class PlainTextProvider(BaseDocumentUnderstandingProvider):
    """纯文本文档理解Provider"""
    
    def __init__(self):
        super().__init__("plaintext")
    
    async def extract_content(self, request: DocumentExtractionRequest) -> DocumentExtractionResult:
        """提取纯文本内容"""
        try:
            # 检测编码
            text = self._decode_text(request.buffer)
            
            # 如果是CSV格式
            if request.file_name.lower().endswith('.csv'):
                return await self._extract_csv_content(request.buffer)
            
            # 如果是HTML格式
            if request.mime and 'html' in request.mime:
                return await self._extract_html_content(request.buffer)
            
            # 普通文本
            return DocumentExtractionResult(
                text=text,
                metadata={
                    "document_type": "plain_text",
                    "encoding": "utf-8",
                    "characters": len(text),
                    "lines": len(text.splitlines())
                }
            )
            
        except Exception as e:
            raise Exception(f"Plain text extraction failed: {str(e)}")
    
    def _decode_text(self, buffer: bytes) -> str:
        """尝试解码文本"""
        # 尝试多种编码
        encodings = ['utf-8', 'gbk', 'gb2312', 'big5', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                return buffer.decode(encoding)
            except UnicodeDecodeError:
                continue
        
        # 如果都失败，使用错误处理
        return buffer.decode('utf-8', errors='ignore')
    
    async def _extract_csv_content(self, buffer: bytes) -> DocumentExtractionResult:
        """提取CSV内容"""
        try:
            text = self._decode_text(buffer)
            
            # 使用csv模块解析
            csv_reader = csv.reader(text.splitlines())
            rows = list(csv_reader)
            
            if not rows:
                return DocumentExtractionResult(
                    text="",
                    metadata={"document_type": "csv", "rows": 0, "columns": 0}
                )
            
            # 格式化CSV内容
            formatted_rows = []
            for i, row in enumerate(rows[:100]):  # 限制行数
                formatted_rows.append("\t".join(row))
            
            formatted_text = "\n".join(formatted_rows)
            if len(rows) > 100:
                formatted_text += f"\n... (共{len(rows)}行，已显示前100行)"
            
            return DocumentExtractionResult(
                text=formatted_text,
                metadata={
                    "document_type": "csv",
                    "total_rows": len(rows),
                    "columns": len(rows[0]) if rows else 0,
                    "displayed_rows": min(100, len(rows))
                }
            )
            
        except Exception as e:
            raise Exception(f"CSV extraction failed: {str(e)}")
    
    async def _extract_html_content(self, buffer: bytes) -> DocumentExtractionResult:
        """提取HTML内容"""
        class HTMLTextExtractor(HTMLParser):
            def __init__(self):
                super().__init__()
                self.text_parts = []
            
            def handle_data(self, data):
                data = data.strip()
                if data:
                    self.text_parts.append(data)
        
        try:
            text = self._decode_text(buffer)
            extractor = HTMLTextExtractor()
            extractor.feed(text)
            
            clean_text = "\n".join(extractor.text_parts)
            
            return DocumentExtractionResult(
                text=clean_text,
                metadata={
                    "document_type": "html",
                    "encoding": "utf-8",
                    "tags_processed": len(extractor.text_parts)
                }
            )
            
        except Exception as e:
            raise Exception(f"HTML extraction failed: {str(e)}")


class DocumentUnderstandingEngine:
    """文档理解引擎"""
    
    def __init__(self):
        self.providers: Dict[str, BaseDocumentUnderstandingProvider] = {}
        self._register_default_providers()
    
    def _register_default_providers(self):
        """注册默认的文档理解Provider"""
        self.providers["plaintext"] = PlainTextProvider()
        
        # 如果PDF库可用，注册PDF Provider
        if PDF_AVAILABLE:
            self.providers["pdf"] = PDFProvider()
        
        # 如果Office库可用，注册Office Provider
        if EXCEL_AVAILABLE:
            self.providers["office"] = OfficeDocumentProvider()
    
    def register_provider(self, provider: BaseDocumentUnderstandingProvider):
        """注册文档理解Provider"""
        self.providers[provider.id] = provider
    
    async def extract_content(
        self, 
        request: DocumentExtractionRequest,
        preferred_provider: Optional[str] = None
    ) -> DocumentExtractionResult:
        """提取文档内容"""
        
        # 根据文件类型自动选择Provider
        if not preferred_provider:
            preferred_provider = self._auto_detect_provider(request)
        
        # 如果指定了首选Provider，优先使用
        if preferred_provider and preferred_provider in self.providers:
            provider = self.providers[preferred_provider]
            try:
                if hasattr(provider, 'extract_content'):
                    return await provider.extract_content(request)
                else:
                    # 对于基础Provider，实现默认提取逻辑
                    return await self._basic_extraction(request, provider.id)
            except Exception as e:
                print(f"Provider {preferred_provider} failed: {e}")
        
        # 尝试所有可用的Provider
        for provider_id, provider in self.providers.items():
            try:
                if hasattr(provider, 'extract_content'):
                    return await provider.extract_content(request)
                else:
                    return await self._basic_extraction(request, provider_id)
            except Exception as e:
                print(f"Provider {provider_id} failed: {e}")
                continue
        
        raise Exception("All document understanding providers failed")
    
    def _auto_detect_provider(self, request: DocumentExtractionRequest) -> str:
        """根据文件类型自动选择Provider"""
        file_name = request.file_name.lower()
        mime_type = request.mime.lower() if request.mime else ""
        
        # PDF文档
        if file_name.endswith('.pdf') or 'pdf' in mime_type:
            return "pdf"
        
        # Office文档
        if any(file_name.endswith(ext) for ext in ['.docx', '.doc', '.xlsx', '.xls', '.pptx', '.ppt']):
            return "office"
        
        # 纯文本
        return "plaintext"
    
    async def _basic_extraction(self, request: DocumentExtractionRequest, provider_id: str) -> DocumentExtractionResult:
        """基础文档提取"""
        try:
            # 简单的文本提取
            text = self._decode_with_fallback(request.buffer)
            
            return DocumentExtractionResult(
                text=text,
                metadata={
                    "provider": provider_id,
                    "file_size": len(request.buffer),
                    "extraction_method": "basic"
                }
            )
            
        except Exception as e:
            raise Exception(f"Basic extraction failed: {str(e)}")
    
    def _decode_with_fallback(self, buffer: bytes) -> str:
        """尝试解码文本（带fallback）"""
        encodings = ['utf-8', 'gbk', 'gb2312', 'big5', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                return buffer.decode(encoding)
            except UnicodeDecodeError:
                continue
        
        return buffer.decode('utf-8', errors='ignore')
    
    async def batch_extract(
        self,
        requests: List[DocumentExtractionRequest],
        preferred_provider: Optional[str] = None,
        max_concurrent: int = 3
    ) -> List[DocumentExtractionResult]:
        """批量提取文档内容"""
        
        semaphore = asyncio.Semaphore(max_concurrent)
        
        async def extract_with_semaphore(request):
            async with semaphore:
                return await self.extract_content(request, preferred_provider)
        
        tasks = [extract_with_semaphore(request) for request in requests]
        results = await asyncio.gather(*tasks, return_exceptions=True)
        
        # 处理异常结果
        processed_results = []
        for i, result in enumerate(results):
            if isinstance(result, Exception):
                processed_results.append(DocumentExtractionResult(
                    text=f"提取失败: {str(result)}",
                    metadata={"error": str(result)}
                ))
            else:
                processed_results.append(result)
        
        return processed_results


# 全局文档理解引擎实例
_document_engine = DocumentUnderstandingEngine()


def register_document_provider(provider: BaseDocumentUnderstandingProvider):
    """注册文档理解Provider"""
    _document_engine.register_provider(provider)


async def extract_document_content(
    request: DocumentExtractionRequest,
    preferred_provider: Optional[str] = None
) -> DocumentExtractionResult:
    """提取文档内容"""
    return await _document_engine.extract_content(request, preferred_provider)


async def batch_extract_documents(
    requests: List[DocumentExtractionRequest],
    preferred_provider: Optional[str] = None,
    max_concurrent: int = 3
) -> List[DocumentExtractionResult]:
    """批量提取文档内容"""
    return await _document_engine.batch_extract(requests, preferred_provider, max_concurrent)


def get_available_document_providers() -> List[str]:
    """获取可用的文档理解Provider列表"""
    return list(_document_engine.providers.keys())